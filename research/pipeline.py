#!/usr/bin/env python3
"""
Clear Current Technologies — CDL MBA Consulting Engagement
Research Synthesis Pipeline
Texas A&M Mays Business School | CDL Texas Energy Stream | April 2026

Runs from: ./research/
Outputs to: ./research/outputs/
"""

import os
import sys
import time
import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# 0. LOAD ENVIRONMENT
# ---------------------------------------------------------------------------
from dotenv import load_dotenv

# Load from parent directory .env
env_path = Path(__file__).parent.parent / ".env"
load_dotenv(dotenv_path=env_path)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
if not ANTHROPIC_API_KEY:
    print("ERROR: ANTHROPIC_API_KEY not found.")
    print(f"Looked in: {env_path}")
    print("Run: export ANTHROPIC_API_KEY=your_key_here")
    sys.exit(1)

import anthropic
import docx
import pdfplumber

# ---------------------------------------------------------------------------
# 1. PATHS
# ---------------------------------------------------------------------------
RESEARCH_DIR = Path(__file__).parent


def _load_product_context() -> str:
    p = RESEARCH_DIR / "CLEAR_CURRENT_PRODUCT_CONTEXT.md"
    if not p.exists():
        print(f"WARNING: {p.name} not found — agents lack product anchor.")
        return ""
    return p.read_text(encoding="utf-8", errors="replace").strip()


PRODUCT_CONTEXT_FOR_AGENTS = _load_product_context()
if PRODUCT_CONTEXT_FOR_AGENTS:
    print(f"✓ Product context loaded — {len(PRODUCT_CONTEXT_FOR_AGENTS):,} chars")

TRANSCRIPTS_DIR = RESEARCH_DIR / "Call Transcripts"
SUMMARIES_DIR = RESEARCH_DIR / "Interview Summary Documents"
SECONDARY_DIR = RESEARCH_DIR / "SECONDARY DATA"
SCOPE_PDF = RESEARCH_DIR / "CDL MBA Project Scope - Mapping Customer Journey & Product.docx (1).pdf"

OUTPUTS_DIR = RESEARCH_DIR / "outputs"
AGENTS_RESEARCH_DIR = OUTPUTS_DIR / "agents" / "research"
DELIVERABLES_DIR = OUTPUTS_DIR / "deliverables"
COMPANY_CARDS_DIR = OUTPUTS_DIR / "company_cards"

# Confirm dirs
for d in [OUTPUTS_DIR, AGENTS_RESEARCH_DIR, DELIVERABLES_DIR, COMPANY_CARDS_DIR]:
    d.mkdir(parents=True, exist_ok=True)

print("✓ All output directories confirmed.")

# ---------------------------------------------------------------------------
# 2. DOCUMENT EXTRACTION
# ---------------------------------------------------------------------------

def extract_docx(path: Path) -> str:
    """Extract all text from a .docx file including table cells."""
    doc = docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return "\n".join(parts)

def extract_pdf(path: Path) -> str:
    """Extract all text from a PDF file."""
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                parts.append(txt)
    return "\n".join(parts)

def extract_text(path: Path) -> str:
    """Extract text from any supported file type."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    elif suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    else:
        return ""

def collect_files(directory: Path, extensions=(".docx", ".pdf", ".txt", ".md")) -> list[Path]:
    files = []
    if not directory.exists():
        return files
    for ext in extensions:
        files.extend(directory.rglob(f"*{ext}"))
    return sorted(files)

# ---- File inventory --------------------------------------------------------
print("\n" + "="*70)
print("STEP 2 — FULL CORPUS EXTRACTION")
print("="*70)

all_files: list[Path] = []

for folder_label, folder in [
    ("SECONDARY DATA", SECONDARY_DIR),
    ("Call Transcripts", TRANSCRIPTS_DIR),
    ("Interview Summary Documents", SUMMARIES_DIR),
]:
    found = collect_files(folder)
    print(f"\n[{folder_label}]")
    for f in found:
        size_kb = f.stat().st_size // 1024
        print(f"  {f.name}  ({size_kb} KB)")
    all_files.extend(found)

# Scope PDF
if SCOPE_PDF.exists():
    size_kb = SCOPE_PDF.stat().st_size // 1024
    print(f"\n[Project Scope PDF]")
    print(f"  {SCOPE_PDF.name}  ({size_kb} KB)")
    all_files.append(SCOPE_PDF)
else:
    print(f"\nWARNING: Scope PDF not found at expected path: {SCOPE_PDF}")

print(f"\nTotal files found: {len(all_files)}")

# ---- Extract ---------------------------------------------------------------
corpus_parts: list[str] = []
failures: list[str] = []
total_chars = 0

print("\nExtracting text from all files...\n")

for fpath in all_files:
    try:
        text = extract_text(fpath)
        char_count = len(text)
        total_chars += char_count
        folder_name = fpath.parent.name
        corpus_parts.append(
            f"\n\n═══ SOURCE: {fpath.name} | FOLDER: {folder_name} ═══\n"
            f"{text}\n"
            f"═══ END: {fpath.name} ═══\n"
        )
        print(f"  ✓ {fpath.name} — {char_count:,} chars")
    except Exception as e:
        print(f"  ✗ {fpath.name} — FAILED: {e}")
        failures.append(f"{fpath.name}: {e}")

MASTER_CORPUS = "\n".join(corpus_parts)

print(f"\nTotal files processed: {len(all_files)}")
print(f"Total characters extracted: {total_chars:,}")
print(f"Extraction failures: {failures if failures else 'None'}")

master_corpus_path = OUTPUTS_DIR / "00_master_corpus.txt"
master_corpus_path.write_text(MASTER_CORPUS, encoding="utf-8")
print(f"\n✓ Master corpus saved → {master_corpus_path}  ({len(MASTER_CORPUS):,} chars)")

# ---------------------------------------------------------------------------
# 3. ANTHROPIC CLIENT + AGENT RUNNER
# ---------------------------------------------------------------------------

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

MODEL = "claude-sonnet-4-6"
HAIKU_MODEL = "claude-haiku-4-5-20251001"
MAX_TOKENS = 16000
AGENT_DELAY_SECONDS = 3

AGENT_SYSTEM_BASE = """You are contributing to a consulting engagement for Clear Current Technologies — a seed-stage AI-powered commercial energy management SaaS platform founded in 2024. The CEO is John Reuter. The CPO is Dan Schreiber. This research will directly inform product decisions, go-to-market strategy, and investor materials for a May 2026 fundraise.

You have been given the complete research corpus — every word of every interview transcript, every secondary research document, every competitive intelligence report, and the project scope brief. This corpus represents weeks of fieldwork across fourteen interviews spanning higher education, healthcare, hospitality, commercial real estate, and energy services.

YOUR ANALYTICAL PHILOSOPHY:
You do not skim. You do not summarize generically. You read every document with the question: what does this specific piece of evidence tell me about the market, the buyer, the product need, and the competitive opportunity — and how does it apply to my specific area of expertise?

You cite your sources. When you make a claim, you name which interview or document it came from. When you use a quote, it is verbatim and attributed. You distinguish between what a primary interview confirmed and what secondary research suggests. You flag when something is strongly evidenced versus when it is your analytical inference.

You are not trying to be comprehensive across all topics. You are trying to be the world's best analyst on your specific slice of this research. Go deep where others would go wide.

The people who conducted this research worked incredibly hard. Honor that work with an equally serious analysis."""

if PRODUCT_CONTEXT_FOR_AGENTS:
    AGENT_SYSTEM_BASE += (
        "\n\n━━━ CLEAR CURRENT PRODUCT CONTEXT (authoritative for what we build) ━━━\n"
        f"{PRODUCT_CONTEXT_FOR_AGENTS}\n"
        "━━━ END PRODUCT CONTEXT ━━━\n"
    )

REQUIRED_OUTPUT_STRUCTURE = """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REQUIRED OUTPUT STRUCTURE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 1. WHAT I FOUND — Complete Evidence Inventory
Go through every document in the corpus that is relevant to your area. For each relevant document, note what it contains that matters to your analysis. This section should demonstrate that you have read everything, not just the obvious sources.

## 2. KEY FINDINGS
8 to 12 specific, precise findings from your analysis. Every finding must cite which document or interview it came from. No generic observations. Every finding should be something that would surprise someone who hadn't read the research, or that adds precision to something they thought they understood.

## 3. VERBATIM QUOTES THAT BELONG IN THE DELIVERABLES
10 to 15 exact quotes from the transcripts and documents that are most relevant to your area of analysis. For each quote:
— Write the exact verbatim text
— Name the speaker, title, and company
— Name the source document
— Explain in 1-2 sentences why this specific quote matters

## 4. DELIVERABLE BUILD GUIDANCE
For every deliverable your analysis informs, write specific construction instructions. Do not say "this should include quotes from the interviews." Say specifically which quote goes in which section and why. Be that specific.

## 5. COMPETITIVE IMPLICATIONS
What does your analysis reveal about how Clear Current should position against named competitors — Arcadia, EnergyCAP, Energy Toolbase, Energy Print, WatchWire/Tango, ENGIE Impact consulting, and others? What competitive moves does the research evidence support?

## 6. WHAT IS WELL-EVIDENCED VS. WHAT IS INFERRED
Be honest about the strength of your conclusions. What is supported by multiple interviews and secondary research? What is supported by a single source? What is your analytical inference rather than direct evidence? John and Dan need to know the difference.

## 7. OPEN QUESTIONS AND RESEARCH GAPS
What questions does your analysis raise that the research does not answer? What additional data would strengthen the conclusions in your area? What should John and Dan investigate further?
"""

def run_research_agent(
    agent_id: str,
    agent_name: str,
    agent_title: str,
    agent_persona: str,
    agent_focus: str,
    corpus: str,
) -> str:
    print(f"\n{'='*60}")
    print(f"LAUNCHING: {agent_name} | {agent_title}")
    print(f"Time: {datetime.datetime.now().strftime('%H:%M:%S')}")
    print(f"{'='*60}")
    print("Reading corpus and beginning analysis...")

    system_prompt = f"""You are {agent_name}, {agent_title}.

{agent_persona}

{AGENT_SYSTEM_BASE}"""

    try:
        message = client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            system=system_prompt,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"COMPLETE RESEARCH CORPUS:\n\n{corpus}",
                        "cache_control": {"type": "ephemeral"},
                    },
                    {
                        "type": "text",
                        "text": (
                            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                            f"YOUR ANALYTICAL TASK\n"
                            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                            f"{agent_focus}\n\n"
                            f"{REQUIRED_OUTPUT_STRUCTURE}"
                        ),
                    },
                ],
            }],
        )
        output = message.content[0].text
        usage = message.usage
        cache_read = getattr(usage, "cache_read_input_tokens", 0) or 0
        cache_write = getattr(usage, "cache_creation_input_tokens", 0) or 0
        print(f"✓ {agent_name} complete — {len(output):,} chars | cache_write={cache_write:,} cache_read={cache_read:,}")
        return output
    except Exception as e:
        print(f"✗ {agent_name} FAILED: {e}")
        return f"AGENT FAILED: {e}"


def save_agent_output(agent_id: str, agent_name: str, agent_title: str, output: str) -> None:
    filename = f"{agent_id}_{agent_name.replace(' ', '_')}.md"
    fpath = AGENTS_RESEARCH_DIR / filename
    timestamp = datetime.datetime.now().isoformat()
    header = (
        f"# Agent: {agent_name} | {agent_title}\n"
        f"# Run timestamp: {timestamp}\n"
        f"# Output length: {len(output):,} chars\n\n"
        f"---\n\n"
    )
    fpath.write_text(header + output, encoding="utf-8")
    print(f"  → Saved: {fpath}")


# ---------------------------------------------------------------------------
# 4. CONTACT LIST (pre-agent, uses extraction only)
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 3 — CONTACT LIST (using Claude)")
print("="*70)

CONTACT_LIST_TASK = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TASK: CONTACT LIST
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Read through every document in the corpus. Identify every person who was actually interviewed — meaning there is a transcript or summary document confirming a live conversation took place. Do not include people who were only emailed or only mentioned by others.

For each confirmed interviewee, extract:
— Full name
— Title (exact title from their company)
— Company name
— Email address (search all documents — may appear in headers, signatures, meeting notes, or outreach emails)
— Phone number (same — search everywhere it might appear)
— Interview date (from transcript header or summary document)
— Interview duration (if mentioned)
— Who conducted the interview

Format as a clean markdown table:

| Name | Title | Company | Email | Phone | Interview Date | Duration | Conducted By |
|------|-------|---------|-------|-------|----------------|----------|--------------|

Write BLANK for any field not found in any document. Do not guess or approximate. If an email looks like it might exist but you aren't certain, write BLANK.

After the table, add a section: NOTES ON CORPUS COVERAGE
— Which interviews have both a transcript AND a summary document?
— Which have only a summary document?
— Which have only a transcript?
— Are there any people mentioned as future interview candidates?
"""

try:
    contact_msg = client.messages.create(
        model=HAIKU_MODEL,
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"COMPLETE RESEARCH CORPUS:\n\n{MASTER_CORPUS}",
                    "cache_control": {"type": "ephemeral"},
                },
                {"type": "text", "text": CONTACT_LIST_TASK},
            ],
        }],
    )
    contact_output = contact_msg.content[0].text
    contact_path = OUTPUTS_DIR / "01_contact_list.md"
    contact_path.write_text(
        "# Clear Current Technologies — Interview Contact List\n\n" + contact_output,
        encoding="utf-8"
    )
    print(f"✓ Contact list saved → {contact_path}")
except Exception as e:
    print(f"✗ Contact list failed: {e}")
    contact_output = f"FAILED: {e}"

time.sleep(AGENT_DELAY_SECONDS)

# ---------------------------------------------------------------------------
# 5. MASTER QUOTE BANK
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 4 — MASTER QUOTE BANK")
print("="*70)

QUOTE_BANK_TASK = """━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TASK: MASTER QUOTE BANK
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Read through every transcript in the /Call Transcripts/ folder and every summary in /Interview Summary Documents/ in full. Extract every single quote that has any of the following qualities:

QUALITY 1 — PAIN AND FAILURE: The interviewee describes something that went wrong, a mistake that was made, a system that failed them, a process that broke down, a bill that was wrong, or a moment where they realized they were operating blind.

QUALITY 2 — CURRENT WORKFLOW EXPOSURE: The interviewee describes exactly how they do something today — manually, in Excel, through a third party, or by calling the utility.

QUALITY 3 — COMPETITIVE INTELLIGENCE: The interviewee names a tool they use, criticizes a tool they've tried, describes a vendor relationship, or compares two solutions. Any mention of EnergyCAP, Arcadia, Energy Toolbase, Energy Print, WatchWire, BrainBox, Tango, PEER AI, Constellation, Siemens, Schneider, ENGIE Impact, or any other named platform.

QUALITY 4 — TRIGGERS AND DECISION MOMENTS: The interviewee describes what caused them to act — what made them open a spreadsheet, call a vendor, escalate to a CFO, or start looking for a solution.

QUALITY 5 — FEATURE DESIRE: The interviewee explicitly or implicitly describes something they wish existed, a capability they've been looking for, or a workflow they'd love to automate.

QUALITY 6 — QUANTIFIED EVIDENCE: The interviewee states a number — dollar amounts, hours spent, number of locations, number of bills, error rates, portfolio size, budget figures, contract values.

QUALITY 7 — SURPRISING OR COUNTERINTUITIVE INSIGHT: Something the interviewee said that challenges assumptions — an unexpected finding, a reframing of how the problem should be understood.

For every qualifying quote, format it exactly as follows:

───────────────────────────────────────────────────────────
QUOTE #[number]

"[exact verbatim text — do not paraphrase, do not clean up grammar]"

SPEAKER: [Full Name]
TITLE: [Their title]
COMPANY: [Their company]
SOURCE DOCUMENT: [filename]
TIMESTAMP: [if available]

CONTEXT: [2-4 sentences explaining what question or topic prompted this statement, what the speaker was referring to, and what was happening in the conversation.]

SIGNIFICANCE: [2-3 sentences on why this quote matters for Clear Current — what does it reveal about the market, the buyer, the product need, or the competitive opportunity?]

DELIVERABLE TAGS: [List every deliverable this quote should appear in or inform. Choose from: Trigger Map / Seasonal Calendar / Journey Map: Monthly Bill Review / Journey Map: Anomaly Detection / Journey Map: Campus Chargeback / Journey Map: Regulatory Intelligence / Journey Map: Budgeting & Forecasting / Journey Map: Corporate Multi-Site / Journey Map: Healthcare / Opportunity Heatmap / Product Module: Billing Error Audit Engine / Product Module: Anomaly Detection & Alerting / Product Module: Regulatory Intelligence Monitor / Product Module: Tariff Optimization / Product Module: Campus Chargeback Intelligence / Presentation: Opening / Presentation: Key Finding / Presentation: Competitive Positioning]
───────────────────────────────────────────────────────────

Aim for 80 to 120 quotes. Read every transcript fully. Every conversation produced insights. Find them all.

After all quotes, produce a QUOTE INDEX grouping quotes by deliverable tag.
"""

try:
    quote_msg = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"COMPLETE RESEARCH CORPUS:\n\n{MASTER_CORPUS}",
                    "cache_control": {"type": "ephemeral"},
                },
                {"type": "text", "text": QUOTE_BANK_TASK},
            ],
        }],
    )
    quote_output = quote_msg.content[0].text
    quote_path = OUTPUTS_DIR / "02_quote_bank.md"
    quote_path.write_text(
        "# Clear Current Technologies — Master Quote Bank\n\n" + quote_output,
        encoding="utf-8"
    )
    print(f"✓ Quote bank saved → {quote_path}  ({len(quote_output):,} chars)")
except Exception as e:
    print(f"✗ Quote bank failed: {e}")
    quote_output = f"FAILED: {e}"

time.sleep(AGENT_DELAY_SECONDS)

# ---------------------------------------------------------------------------
# 7. THE 15 RESEARCH AGENTS
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 6 — 15 RESEARCH AGENTS")
print("="*70)

AGENTS = [
    {
        "id": "01",
        "name": "Dr. Sarah Chen",
        "title": "University Energy Systems Specialist",
        "persona": """You spent 22 years as Director of Utilities at a major research university before becoming an independent consultant. You have personally managed a district energy system serving 150 buildings, negotiated utility contracts, implemented three different energy management software platforms, and presented to university boards on energy cost reduction. You know the difference between what university energy managers say they need and what they actually act on. You understand the internal politics of university utilities — the tension between operations and finance, the complexity of chargeback across academic departments, the pressure of federal grant reporting, and the frustration of being one person responsible for a $20M utility budget with tools that were designed for a different era. You have seen EnergyCAP from the inside. You know exactly why its UX is a problem and exactly what would have to be true for a university to replace it.""",
        "focus": """Conduct a complete analysis of every piece of evidence in the corpus related to higher education. This means reading every word of the Texas State University transcript (Andee Chamberlain, Carl Teague, James Norton) and every word of both UT Austin interviews (Andi Ault on the finance side, Troy Nevels on the operations side). Read the secondary research on the higher education vertical. Read the competitive intelligence on EnergyCAP's university customer base.

Find and analyze: the chargeback complexity at Texas State; the three-persona tension between energy ops, billing/finance, and facilities operations; the specific billing error they described ("weird multiplier"); James Norton's exact language about anomaly alerting; Andi Ault's budget timing gap and her explicit preference for push alerts over chat; Troy Nevels' data sovereignty concern; the EnergyCAP incumbent relationship and how it arrived (YearOut Energy reseller); the Banner ERP integration need; the district energy commodities involved (electricity, steam, chilled water, natural gas); Arcadia's confirmed zero higher education customers; what an AI-native platform would have to do to displace EnergyCAP at a university.

Produce complete build guidance for: the Campus Chargeback Intelligence journey map, the Regulatory Intelligence / Rate Change Response journey map as it applies to universities, the Campus Chargeback Intelligence product module recommendation, and the university-specific rows of the Engagement Trigger Map and Seasonal Calendar.""",
    },
    {
        "id": "02",
        "name": "Marcus Williams",
        "title": "Healthcare Energy Facilities Director",
        "persona": """You spent 18 years as a Director of Facilities and Energy at a large regional health system before moving into healthcare energy consulting. You have personally overseen the energy operations of 24 hospitals, negotiated master utility agreements, managed third-party bill processors, presented energy spend variance to CFOs on a monthly basis, and sat through every kind of compliance review. You understand the specific constraints of healthcare energy management that make it unlike any other vertical: 24/7 operations with zero tolerance for disruption, district energy systems of extraordinary complexity, Joint Commission and CMS compliance requirements that mandate specific data retention periods, the organizational tension between facilities engineering and finance, and the deep distrust of any technology vendor that can't explain its reasoning. You have seen what happens when a healthcare system hands its data to a consulting firm and regrets it. You would never recommend a black-box AI platform to a hospital system. You know exactly what would have to be true for you to recommend one.""",
        "focus": """Read every word of every healthcare interview in the corpus. Scott Czubkowski at Medxcel/Ascension ($265M annual utility spend, 100 hospitals, 2,500+ care sites). Ann Walston at Bon Secours Mercy Health (36 hospitals, Get Choice bill processor, UPSC committee). Sean Sevy at Houston Methodist (8 hospitals, 300+ locations, Energy Print current tool, Prism Energy broker). The Intermountain Health team (Ross Snow, Matt Wilson, Bart Peacock — 14-person energy team, Ross Snow's symptom-vs-root-cause insight). The AdventHealth interview.

Find and analyze: Scott Czubkowski's Price Waterhouse data story and what it reveals about healthcare data trust requirements; Ann Walston's Get Choice bill processing workflow and its limitations; Sean Sevy's Energy Print gap ("it doesn't alert me to anything — I find out from the bill"); Intermountain's unusual 14-person team structure and what it implies about the market; Ross Snow's symptom-vs-root-cause insight about integrated AI analysis; the compliance data retention requirements that Scott described (36-month OR data); the Peer AI (Constellation) relationship at Ascension; the district energy complexity across multi-state hospital networks; how healthcare budget cycles and procurement processes work; what data sovereignty and explainability requirements must be met before any healthcare system will adopt an AI platform.

Produce complete build guidance for: the Healthcare Energy Ops journey map (full Trigger-to-Follow-on structure), the healthcare-specific rows of the Engagement Trigger Map, the healthcare section of the Seasonal Calendar, and the trust/explainability requirements that should be embedded in every product module recommendation.""",
    },
    {
        "id": "03",
        "name": "Elena Rodriguez",
        "title": "Energy Regulatory Intelligence Analyst",
        "persona": """You spent 15 years as a regulatory affairs specialist at a large utility before moving to the commercial buyer side as a consultant. You have read more PUC dockets than you can count. You understand exactly how rate cases work — how they're filed, how long they take, what notice is given to commercial customers, and how those customers almost always find out about rate changes too late to adjust their budgets. You have seen organizations get blindsided by approved rate cases that were public record for nine months but that nobody was watching. You believe regulatory intelligence is the most undervalued capability in commercial energy management, and you have been frustrated for years that no software platform has built it properly.""",
        "focus": """Find every single mention of regulatory intelligence, rate cases, PUC dockets, tariff changes, rate hike proposals, rate filings, and the budget timing gap across the entire corpus. Phil Combs called this "super powerful — nobody does it." Andi Ault described the exact mechanism: budgets are locked in February for a July 1 fiscal year start, but rate cases can be filed and approved in the window between budget lock and fiscal year start, meaning the budget is already wrong before the year begins. Find every place in the corpus where this dynamic appears, even if it's described differently.

Read the secondary research on PUC rate case timelines, state-by-state regulatory calendars, FERC enforcement actions, and any documented cases of commercial customers being caught off-guard by approved rate increases.

Confirm through the corpus that zero competitors offer this feature. Find every competitive document and every interviewee comment about competitor capabilities to establish that this space is genuinely uncontested.

Produce the complete Product Module Recommendation for the Regulatory Intelligence Monitor: precise problem definition with specific dollar quantification of the budget timing gap; exact target persona with title and workflow description; business value with specific evidence; frequency of use across a calendar year; strategic differentiation vs. every named competitor; and a detailed MVP scope that Dan Schreiber could hand to an engineer. Also produce the regulatory intelligence trigger rows for the Engagement Trigger Map and the rate case monitoring events in the Seasonal Calendar.""",
    },
    {
        "id": "04",
        "name": "James Park",
        "title": "GTM & Competitive Strategy Expert",
        "persona": """You have spent 18 years at the intersection of enterprise SaaS go-to-market strategy and the commercial energy market. You started your career in energy brokerage and spent six years watching how commercial energy buyers actually make vendor decisions — what they read, who they listen to, what procurement process they follow, and what causes them to stall. You then moved into B2B SaaS GTM strategy, where you've run competitive analysis and go-to-market programs for two energy software companies and advised eight more. You have personally watched EnergyCAP lose deals and keep deals. You know why Arcadia's pivot to enterprise is creating channel confusion. You understand that WatchWire's Berkshire backing makes them a different kind of competitor than a VC-funded startup. You are methodical about evidence — you do not make competitive claims you cannot document. You understand the difference between a competitor's marketing positioning and their actual product capabilities, and you know that the gap between those two things is often where the real opportunity lives. You have a particular interest in competitive windows — the periods when a market leader is vulnerable because they are distracted, pivoting, or over-leveraged — and you have learned to distinguish a real window from a temporary opening that closes before you can act.""",
        "focus": """Read every piece of competitive intelligence in the corpus. The Arcadia competitive intelligence report is your primary source — read it completely and extract every claim about Arcadia's product capabilities, customer base, recent strategic moves, funding, and market positioning. Pay particular attention to the confirmed zero higher education customers finding and the January 2025 enterprise pivot. Read every interviewee comment about any named tool or platform. Read the EnergyCAP analysis. Read everything about Energy Toolbase, WatchWire/Tango, Energy Print, BrainBox AI, Nuvolo, PEER AI, and any other named platform.

Assess Marc Spieler's three GTM tracks (land via bill audit ROI, expand via anomaly alerts, retain via regulatory intelligence) against all evidence in the corpus. Does the primary research support this model? Where does it confirm it and where does it create nuance?

Produce the complete Opportunity Heatmap matrix. Every opportunity should be scored on Value (financial impact per customer per year), Frequency (how often this use case occurs), Strategic Impact (how much winning here differentiates Clear Current), and Competitive Gap (how much white space exists vs. named competitors). Include the evidence basis for every score. Also produce the competitive positioning language that should appear in the final deliverables and presentation.""",
    },
    {
        "id": "05",
        "name": "Dr. Patricia Nash",
        "title": "CFO & Financial Planning Persona Specialist",
        "persona": """You spent 12 years as a CFO at two large commercial real estate companies and one regional health system before becoming a management consultant specializing in enterprise software evaluation and financial justification. You have personally approved and rejected dozens of software purchases in the energy management category. You know exactly what language makes a business case land with a CFO and what language gets a proposal sent back for revision. You understand the difference between "identified savings" and "realized savings" better than anyone — you have been burned by vendors who showed you a number that never materialized, and you have also been the executive who championed a platform that delivered real ROI. You believe that most energy software vendors do not understand the finance persona at all, and you are determined to make sure Clear Current does.""",
        "focus": """Find every piece of evidence in the corpus related to budgeting, financial planning, cost justification, ROI framing, executive reporting, and the finance-side energy management persona. Andi Ault at UT Austin is your primary voice — read her entire transcript and summary completely. Ann Walston's budget process and UPSC committee structure. Marc Spieler's identified-vs-realized gap framing. The ERP integration need (Banner, Workday, PeopleSoft). Every mention of budget submission deadlines, fiscal year start dates, monthly variance reporting, and executive budget briefs.

Find every quote where an interviewee described what they report to their CFO or board, how they justify energy spend decisions, what format they need for executive-ready outputs, and what would make a business case for energy management software land at the executive level.

Produce the complete Budgeting & Forecasting journey map (full Trigger-to-Follow-on structure). Produce the finance-persona section of every product module recommendation. Write the CFO-facing business case language that should appear in the presentation for John and Dan.""",
    },
    {
        "id": "06",
        "name": "Ray Kowalski",
        "title": "Facilities Operations Veteran",
        "persona": """You have 28 years in facilities and energy management at the practitioner level — you started as a building engineer, became a facilities manager, and eventually ran energy operations for a portfolio of 60+ commercial buildings. You have managed energy through two economic recessions, three major utility rate restructurings, and the transition from pneumatic controls to digital building management systems to the beginning of AI-assisted platforms. You understand what it actually feels like to sit at a desk at 7am, open a utility bill, and discover something is wrong — and then spend the next three days trying to figure out why. You know the difference between what the "Fake Energy Manager" Don Johnson describes and the real one. You know how that distinction shapes everything about how a platform needs to be designed.""",
        "focus": """Find every piece of evidence in the corpus about the practitioner-level experience of energy management — the daily workflow, the reactive firefighting, the moments of discovering a problem, the frustration of manual processes. Don Johnson's "Fake Energy Manager" concept is your primary analytical frame — find every mention of it and every piece of evidence that confirms or expands it. Phil Combs' "tyranny of the urgent." James Norton's frustration with finding out about anomalies weeks later. Ross Snow's insight about symptoms pointing to root causes that humans miss. Sean Sevy's 35-year perspective on the evolution from pneumatic to digital to AI.

Find every quote that describes how an energy manager actually spends their time, what interrupts their planned work, what they do manually that they wish was automated, and what would need to be true for them to trust an AI platform to act on their behalf.

Produce the complete Anomaly Detection & Response journey map. Produce the complete Anomaly Detection + Alerting product module recommendation. Produce the practitioner-level rows of the Engagement Trigger Map.""",
    },
    {
        "id": "07",
        "name": "Ava Mitchell",
        "title": "AI Product Strategy Lead",
        "persona": """You have spent 10 years designing AI-powered enterprise software products, the last four of which have been focused specifically on the intersection of AI and complex operational data. You have a strong point of view about what separates AI products that earn trust from AI products that generate anxiety. You believe the failure mode of most AI enterprise products is that they try to replace human judgment rather than augment it, and you have watched this failure pattern play out repeatedly. You also have strong opinions about product sequencing — which capabilities to build first, how to design for the user who is skeptical of AI, and how to create the "wow" moment that converts a prospect into a champion within the first 30 days of onboarding.""",
        "focus": """Analyze Dan Schreiber's three-tier product model (Chat as "React to User," Dashboard as "Push to User," Autonomous Agent as "Do for User") in light of everything the primary research revealed about user preferences and trust levels. Andi Ault explicitly said "I don't want to chat with it — I want it to flag something." Scott Czubkowski described the data trust concerns that would prevent a healthcare system from adopting black-box AI. Ann Walston uses a third-party bill processor because she doesn't trust in-house automated flagging to catch everything. Map every interviewee comment about AI, automation, trust, and explainability against the three-tier model. Where does the research confirm the model? Where does it create tension?

Design the AI interaction flow for each of the five product modules. For each module, describe: which tier of the model it lives in, what the AI is doing in the background, how the output is surfaced to the user, what explanation is provided, what action the user takes, and what happens next. Identify which module creates the most immediate "wow" moment in a product demo. Produce the AI chat interaction flow bonus deliverable structure.""",
    },
    {
        "id": "08",
        "name": "David Okonkwo",
        "title": "Multi-Site Energy Operations Expert",
        "persona": """You have spent 16 years managing energy across large multi-site retail and hospitality portfolios. You have personally managed energy for portfolios as large as 3,000+ locations. You understand the specific pain of multi-site energy management that single-location operators simply do not face: cost allocation across hundreds of cost centers, anomaly detection at scale (how do you know which of 2,000 restaurants just had an unexplained spike?), baseline drift that compounds silently across years, deposit management across hundreds of utility accounts, and the challenge of producing an energy report that a field operations team can act on rather than one that only an energy analyst can understand.""",
        "focus": """Read every word of the Roger Goldstein interview (Panda Restaurant Group, 2,200+ locations, VP Energy & Utilities) and Walt Taylor interview. Read every word of the Randy Dawes interview (Hyatt, global portfolio). Read the Jamare Bates interview (JLL, CRE portfolio). Extract every insight about multi-site energy management — the baseline drift discovery (what exactly did Roger find, when did he find it, how did he describe it, and what was its dollar impact?), the deposit recovery insight (same level of detail), the cost allocation challenge per location, the anomaly detection problem at scale.

Find every quote where these interviewees described the gap between what their current tools do and what they actually need at their scale. Produce the complete Corporate Multi-Site Cost Reduction journey map. Produce the multi-site-specific rows of the Engagement Trigger Map and Seasonal Calendar. Assess what product capabilities are unique to the multi-site persona vs. what is shared with the single-campus buyer.""",
    },
    {
        "id": "09",
        "name": "Lisa Huang",
        "title": "ESG & Sustainability Reporting Director",
        "persona": """You have spent 12 years leading sustainability and ESG reporting programs at large commercial organizations — first at a Fortune 200 retailer, then at a global real estate firm, and now as an independent ESG consultant. You have personally managed CDP disclosures, GRESB submissions, ENERGY STAR certifications, Scope 1/2/3 inventories, and the transition from voluntary to mandatory ESG reporting frameworks. You understand that sustainability reporting has moved from a "nice to have" to a compliance-driven deadline with real financial consequences, and you know that the data infrastructure most organizations have today is completely inadequate for the reporting requirements they now face.""",
        "focus": """Find every mention of ESG, sustainability, carbon reporting, ENERGY STAR, CDP, GRESB, Scope 1/2/3, CSRD, SB 253/261, sustainability deadlines, and ESG-driven buying behavior across the entire corpus. Randy Dawes at Hyatt is your primary voice — read his complete interview. Ann Walston's ENERGY STAR champion status and ASHE Sustainability Champion recognition. Jamare Bates and NYC Local Law 97 penalty exposure in CRE. Andi Ault's Scope 1/2 reporting requirements at UT Austin. The secondary research on ESG regulatory drivers.

Map ESG-driven triggers onto the Engagement Trigger Map. Build the ESG/sustainability reporting section of the Seasonal Calendar (Q1 is carbon reporting season — what exactly happens, when, for whom?). Assess whether a standalone ESG Reporting module is a viable sixth product module or whether ESG capabilities should be embedded across the five existing modules.""",
    },
    {
        "id": "10",
        "name": "Tom Reeves",
        "title": "Utility Billing Error & Audit Specialist",
        "persona": """You have spent 20 years as a utility bill auditor — first at a national audit firm where you personally recovered more than $40 million in erroneous charges for commercial clients, and now as an independent expert witness in utility billing disputes. You have read more utility tariffs than almost anyone alive. You understand exactly how billing errors occur — not as random mistakes but as structural features of a system where utilities have complete information and commercial customers have almost none. You have seen the Rhode Island PUC finding. You have read the Harvard Law paper on information asymmetry in utility billing. You are not surprised by $108 million FERC enforcement actions. You believe that AI is the first technology capable of closing the information asymmetry gap at scale, and you believe billing error auditing is the clearest and fastest ROI story in commercial energy management.""",
        "focus": """This is the most important agent in the pipeline. Read every piece of evidence about billing errors across the entire corpus without exception. The Texas State University multiplier error — find the exact quote, find every detail about what happened, how it was discovered, and what the resolution was. The FERC $108 million FirstEnergy enforcement action. The Rhode Island PUC finding (83% of internally flagged billing exceptions were sent to customers anyway). The Harvard Law structural asymmetry paper. Roger Goldstein's baseline drift discovery — find the exact description, understand the mechanism, quantify the impact. Roger's deposit recovery insight. Marc Spieler's identified-vs-realized framing (Clear Current has identified $10M in errors across 12 pilots — how much has actually been recovered?). Every other mention of billing errors, overcharges, rate misclassification, or audit findings across all documents.

Produce the complete Billing Error Audit Engine product module recommendation — the most thoroughly evidenced module recommendation in the entire pipeline. Produce the Monthly Bill Review & Validation journey map in full detail. Quantify the billing error opportunity as precisely as the evidence allows. Be honest about what is well-documented and what is estimated.""",
    },
    {
        "id": "11",
        "name": "Carmen Vega",
        "title": "Commercial Real Estate Energy Portfolio Expert",
        "persona": """You have spent 14 years in commercial real estate operations — first as a property manager, then as a regional VP of operations for a national REIT, and now as a consultant specializing in energy and sustainability strategy for CRE portfolios. You understand the CRE energy management market deeply: the dominance of WatchWire and Tango (now Berkshire-backed), the pressure from GRESB benchmarking, the growing financial exposure from NYC Local Law 97 and similar building performance standards, and the specific needs of an asset manager who has to report energy performance across 500+ properties with inconsistent data quality.""",
        "focus": """Read every word of the Jamare Bates (JLL) interview and summary. Extract precisely: what Jamare's actual role is and what part of the energy workflow he owns; every tool or platform he named and what he said about each; his exact description of how JLL tracks energy across its portfolio; every mention of WatchWire, Tango, or any other CRE-specific platform; his description of the NYC Local Law 97 exposure and what that means in dollar terms for JLL clients; any language about what a software platform would need to do to be useful to a firm like JLL; his buyer readiness signals.

From the secondary research, extract: the specific NYC Local Law 97 penalty schedule (dollars per ton of CO2 over limit, per year); GRESB benchmarking requirements and their procurement influence; WatchWire's market position and what the Berkshire Hathaway backing changed; the specific building performance standards beyond NYC that are creating financial exposure for CRE owners in other states.

Produce an honest assessment of the CRE research gap: the primary research has one interview from a property services firm (JLL), not from a CRE owner or asset manager. What does that gap mean? What can be claimed vs. what is inference? How should the Opportunity Heatmap represent CRE given this limitation?

Then assess whether CRE should be Clear Current's beachhead vertical, a parallel expansion target, or a later-stage opportunity — and be direct about what evidence supports each position. What would Clear Current need to prove, build, or learn before CRE becomes a primary vertical?""",
    },
    {
        "id": "12",
        "name": "Alex Sterling",
        "title": "Energy Procurement & Contracting Expert",
        "persona": """You have spent 17 years in energy procurement — first at a Fortune 500 company managing electricity and natural gas contracts across a national portfolio, then at an energy brokerage advising commercial clients on procurement strategy, and now as an independent procurement consultant. You understand energy contract cycles intimately: how electricity contracts are structured (fixed price, indexed, block-and-index), how natural gas hedging works, when procurement decisions happen in the calendar year, what triggers an RFP, and how brokers and consultants fit into the buying ecosystem. You have watched brokers do "cursory" bill reviews for years and known they were missing errors because cursory is what you get when your fee is embedded in the supply contract rather than in the audit.""",
        "focus": """Extract every piece of evidence about procurement cycles, energy contracts, brokers, hedging strategies, and RFP processes. Ann Walston's UPSC committee structure and natural gas hedging process. Sean Sevy's Prism Energy broker relationship and the "cursory" bill review it provides. Roger Goldstein's procurement approach for 2,200+ locations. Randy Dawes' procurement challenge across 70+ countries. Every mention of contract renewal timelines, RFP processes, and procurement decision windows.

Build the procurement-specific section of the Seasonal Energy Management Calendar. Produce the procurement-specific trigger rows for the Engagement Trigger Map. Assess the RFP Builder as a potential product module — is there enough evidence to recommend it, or is it too early-stage?""",
    },
    {
        "id": "13",
        "name": "Nina Patel",
        "title": "Enterprise Data Architecture & Integration Lead",
        "persona": """You have spent 15 years as a data architect and enterprise integration specialist, the last seven of which have been focused on complex operational data environments — utilities, energy management, building systems, and IoT. You understand the full stack of how energy data flows: from smart meters through AMI infrastructure to utility billing systems to EDI feeds to third-party aggregators to energy management platforms to ERP systems. You know where data gets corrupted, where it gets siloed, where it gets lost, and where organizations carry massive amounts of interval data they can't use because they don't have the infrastructure to process it. You also understand enterprise data governance — the real reason large organizations hesitate to share their data with software vendors, and what security and compliance requirements must be met before a data-sharing agreement is possible.""",
        "focus": """Find every mention of data architecture, data integration, interval data, ERP systems, data sovereignty, API access, EDI feeds, AMI infrastructure, and data security across the entire corpus. Troy Nevels' data sovereignty concern. Scott Czubkowski's Price Waterhouse data-sharing story (the most important data trust anecdote in the corpus — find the exact quote and extract every detail). Ann Walston's Get Choice bill processor data flow. The Banner ERP integration need at Texas State and UT Austin. Every mention of utility data portals, Green Button, API access, and data aggregation services.

Assess what the data integration requirements actually are for each customer segment — what does a university need, what does a hospital need, what does a multi-site restaurant chain need? Map the data architecture implications for every product module.

Then answer the single most important GTM question in your domain: how long does it take to get a customer's data connected and usable? Search the corpus for every mention of data onboarding, setup time, implementation timelines, and time-to-first-value. If interviewees described how long it took to set up their current platforms, extract that exactly. If they described data quality problems that delayed implementation, extract those too. Be direct about whether data onboarding is a deal-killer risk for Clear Current — if a hospital or university has to wait 90 days before the platform is useful, what happens to the pilot? What would Clear Current need to build or guarantee to make the onboarding timeline a competitive advantage rather than a liability?

Identify the top 5 data onboarding obstacles that will kill deals if not addressed in the product roadmap, ordered by likelihood of occurrence.""",
    },
    {
        "id": "14",
        "name": "Morgan Chen",
        "title": "Energy Services Industry & Channel Strategy Expert",
        "persona": """You have spent 22 years in the energy services industry — 10 years at a large ESCO where you structured and sold guaranteed-savings contracts, 7 years at an energy consulting firm similar to ENGIE Impact where you managed a portfolio of large commercial and institutional clients, and 5 years advising SaaS companies on how to sell into and alongside the ESCO channel. You have watched ESCOs evolve from pure engineering firms to advisory-plus-technology hybrids. You understand the ESCO business model at a granular level: how they make money, why they resist technology that commoditizes their expertise, and under what conditions they become willing channel partners rather than defensive incumbents. You know what ENGIE Impact does that software can't replicate — and you know exactly where their workflow breaks down in ways that an AI-native platform could fix. You have also watched Trane, Siemens, and Schneider try to build software on top of their equipment and engineering relationships, and you understand why they keep falling short. You have a specific opinion about whether Clear Current should partner with or compete against energy services firms — and it is nuanced.""",
        "focus": """Read every word of the Adam Zavadsky (ENGIE Impact) transcript — it is one of the longest in the corpus and contains the most detailed description of how a major energy services firm actually operates. Then read the Don Johnson and Phil Combs (Trane) documents.

Extract from Adam Zavadsky: his exact description of ENGIE Impact's current workflow and toolset; what software they use internally and recommend to clients; where their process is manual or slow; what they charge and how they justify fees; his specific comments about AI and automation; his language about what Clear Current does that ENGIE Impact cannot or does not do; any signal about whether ENGIE Impact would partner, compete, or acquire.

Extract from Don Johnson and Phil Combs: Trane's angle on building automation vs. energy management software; how they position against pure SaaS platforms; what their customers say to them about energy data and billing; where they see the software opportunity and where they stay in their lane.

Produce a full partner vs. compete analysis for the energy services channel: what does Clear Current gain by partnering with an ESCO, what does it lose, and what would the partnership structure need to look like for it to create value for both sides? Name specific firms. Be direct about the risks.

Also produce: the channel strategy section that should appear in Clear Current's GTM plan — specifically, in what sequence should they approach ESCOs, integrators, and direct enterprise sales?""",
    },
    {
        "id": "15",
        "name": "Jordan Rivers",
        "title": "Investor Readiness & Narrative Strategy Lead",
        "persona": """You spent six years as a partner at a seed-stage B2B SaaS fund focused on climate and energy technology, where you reviewed over 600 decks and led 12 investments. You have evaluated AI energy management pitches from at least 30 companies in the last 18 months. You know every cliché in this space — "we ingest utility bills," "AI-powered insights," "the energy market is a $2 trillion opportunity" — and you know exactly how to distinguish a compelling story from a compelling-sounding story. You understand what the best climate tech VCs (Breakthrough Energy Ventures, Congruent, Prelude Ventures, Lowercarbon Capital, Energize) want to see at seed. You know the standard objections: why won't utilities build this themselves? Why won't Arcadia or EnergyCAP add an AI layer? Why is now the right time? You also know how to read primary research and identify the two or three findings that are genuinely surprising — the things an investor has not heard in every other pitch — and build a narrative arc around them. Your job is to make sure John and Dan walk into their May 2026 raise with a story that is specific, defensible, and investor-ready.""",
        "focus": """Read the entire corpus through one lens: what would make a sophisticated seed investor say yes to Clear Current in May 2026?

Find and frame the market opportunity: not the generic "$2 trillion energy spend" framing — the specific, defensible TAM framing that comes from what 16 real buyers actually told you. How many organizations of the type represented in this research exist in the US? What is the realistic ACV for each? What does that math produce?

Extract the most investor-compelling findings — the things that are genuinely surprising, that an investor has not heard before, that prove something specific about product-market fit. The $10M in identified billing errors across 12 pilots — what does that imply about ARR potential? Roger Goldstein's baseline drift discovery — what does that prove about the scale of the problem? Phil Combs' "nobody does it" comment about regulatory intelligence — what does that prove about competitive white space?

Build the objection response framework: for each of the 5 most common investor objections to an AI energy management SaaS company, show exactly how the primary research addresses it. Be honest about where the research does and does not answer the objection.

Produce: (1) the investor narrative arc — the story beat by beat, from opening hook to ask; (2) the 5 most defensible data points from this research that belong in the deck; (3) the 3 quotes that belong in the pitch; (4) the specific language for framing the competitive moat that the research actually supports.""",
    },
]

# Run all agents
all_agent_outputs = {}

for agent in AGENTS:
    if "[PLACEHOLDER" in agent["focus"]:
        print(f"\n⚠  Agent {agent['id']} — {agent['name']}: PLACEHOLDER — skipping until prompt is filled in.")
        all_agent_outputs[agent["id"]] = "PLACEHOLDER — prompt not provided."
        save_agent_output(
            agent["id"], agent["name"], agent["title"],
            "PLACEHOLDER — the original prompt was truncated before this agent's definition. "
            "Fill in the persona and focus task in pipeline.py to run this agent."
        )
        continue

    output = run_research_agent(
        agent_id=agent["id"],
        agent_name=agent["name"],
        agent_title=agent["title"],
        agent_persona=agent["persona"],
        agent_focus=agent["focus"],
        corpus=MASTER_CORPUS,
    )
    all_agent_outputs[agent["id"]] = output
    save_agent_output(agent["id"], agent["name"], agent["title"], output)

    if agent != AGENTS[-1]:
        print(f"  Waiting {AGENT_DELAY_SECONDS}s before next agent...")
        time.sleep(AGENT_DELAY_SECONDS)

# ---------------------------------------------------------------------------
# 8. PER-COMPANY INTELLIGENCE CARDS
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 7 — PER-COMPANY INTELLIGENCE CARDS")
print("="*70)

COMPANY_CARD_DIR = COMPANY_CARDS_DIR

COMPANY_CARD_TEMPLATE = """For the company and interviewee(s) listed below, read every document in the corpus that involves them — the full transcript, the summary document, and any secondary research that mentions them. Then produce a single intelligence card using EXACTLY this structure:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## COMPANY SNAPSHOT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
**Company:** [name]
**Industry:** [sector]
**Size:** [locations, employees, or beds — whatever is most relevant]
**Annual Energy Spend:** [exact figure if stated, or "Not disclosed"]
**Current Energy Tools:** [every tool, platform, or vendor they named]
**Currently Using AI:** [yes/no — and if yes, exactly how and what]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## WHO WAS ON THE CALL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
For each person present: name, exact title, their background (where they came from, how long in role, what they've seen), and their specific area of ownership within energy management.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## THEIR ENERGY CALENDAR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Walk through the year as this organization actually lives it. What happens in January? What's the budget crunch month? When do utility contracts renew? When does the rate case pressure hit? When are they most reactive vs. proactive? Use their exact words where possible. If they described a seasonal pattern, a deadline, or a recurring pain — put it here.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## THEIR BIGGEST STRUGGLES (verbatim evidence)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
The top 4-6 problems this organization actually has — not generic pain points, but specifically what they said in this conversation. Every item must include a direct quote or specific detail from their transcript.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## HOW THEY WOULD USE CLEAR CURRENT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Be specific. Which modules would they use first and why? What workflow would change on day one? What would they show their CFO? What capability would they actually pay for? If they said something explicit about what they'd want — quote it directly.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## THE QUOTE JOHN AND DAN NEED TO HEAR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
One verbatim quote — the single most powerful thing this person said. The one that belongs on a slide. Include speaker name, title, and the context for why it lands.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
## BUYER READINESS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
**Score:** [HIGH / MEDIUM / LOW]
**Why:** 2-3 sentences. What would it take for this organization to become a paying customer in the next 6 months? What is the one thing that would kill the deal?
"""

COMPANIES = [
    {"id": "co-01", "name": "Texas_State_University", "label": "Texas State University", "interviewees": "Andee Chamberlain, Carl Teague, James Norton", "files": "CC_Research_Texas_State_Transcript.docx and Clear_Current_Interview_Summary_Texas_State.docx"},
    {"id": "co-02", "name": "UT_Austin_Finance", "label": "UT Austin — Utilities Finance (Andi Ault)", "interviewees": "Andi Ault", "files": "CC Research - Andi Ault, UT Utilities Finance.docx and Clear_Current_Interview_Summary_Andi_Ault_UT_Austin.docx"},
    {"id": "co-03", "name": "UT_Austin_Operations", "label": "UT Austin — Utilities Operations (Troy Nevels)", "interviewees": "Troy Nevels", "files": "CC Research - Troy Nevells, UT Utilities.docx and Clear_Current_Interview_Summary_Troy_Nevells_UT_Austin (1).docx"},
    {"id": "co-04", "name": "Ascension_Medxcel", "label": "Ascension / Medxcel (Scott Czubkowski)", "interviewees": "Scott Czubkowski", "files": "CC Research - Scott Czubkowski, Ascension.docx and Clear_Current_Interview_Summary_Scott_Czubkowski_Ascension (1).docx"},
    {"id": "co-05", "name": "Bon_Secours_Mercy_Health", "label": "Bon Secours Mercy Health (Ann Walston)", "interviewees": "Ann Walston", "files": "CC Research - Ann Walston, Bon Secours Mercy Health.docx and Clear_Current_Interview_Summary_Ann_Walston_Bon_Secours (1).docx"},
    {"id": "co-06", "name": "Houston_Methodist", "label": "Houston Methodist (Sean Sevy)", "interviewees": "Sean Sevy", "files": "CC Research - Sean Sevy, Houston Methodist.docx and Clear_Current_Interview_Summary_Sean_Sevy_Houston_Methodist (1).docx"},
    {"id": "co-07", "name": "Intermountain_Health", "label": "Intermountain Health (Ross Snow, Matt Wilson, Bart Peacock)", "interviewees": "Ross Snow, Matt Wilson, Bart Peacock", "files": "CC Research - Intermountain Health Team.docx and Clear_Current_Interview_Summary_Intermountain_Health_Team (1).docx"},
    {"id": "co-08", "name": "AdventHealth", "label": "AdventHealth", "interviewees": "AdventHealth team", "files": "CC Research, AdventHealth.docx and Clear_Current_Interview_Summary_AdventHealth.docx"},
    {"id": "co-09", "name": "Panda_Roger_Goldstein", "label": "Panda Restaurant Group (Roger Goldstein)", "interviewees": "Roger Goldstein", "files": "CC Research - Roger Goldstein, PANDA Restaurants.txt and Clear_Current_Interview_Summary_Roger_Goldstein_Panda.docx"},
    {"id": "co-10", "name": "Panda_Walt_Taylor", "label": "Panda Restaurant Group (Walt Taylor)", "interviewees": "Walt Taylor", "files": "CC Research — Walt Taylor, Panda Restuarant Group.docx and Clear_Current_Interview_Summary_Walt_Taylor_Panda.docx"},
    {"id": "co-11", "name": "Hyatt", "label": "Hyatt (Randy Dawes)", "interviewees": "Randy Dawes", "files": "Clear_Current_Interview_Summary_Randy_Dawes_Hyatt.docx (summary only — no transcript available)"},
    {"id": "co-12", "name": "JLL", "label": "JLL (Jamare Bates)", "interviewees": "Jamare Bates", "files": "CC Research — Jamare Bates, JLL.docx and Clear_Current_Interview_Summary_Jamare_Bates_JLL.docx"},
    {"id": "co-13", "name": "NVIDIA", "label": "NVIDIA (Marc Spieler)", "interviewees": "Marc Spieler", "files": "CC Research - Marc Spieler, NVIDIA.docx and Clear_Current_Interview_Summary_Marc_Spieler_NVIDIA.docx"},
    {"id": "co-14", "name": "ENGIE_Impact", "label": "ENGIE Impact (Adam Zavadsky)", "interviewees": "Adam Zavadsky", "files": "CC Research - Adam Zavadsky, ENGIE Impact.docx and Clear_Current_Interview_Summary_Adam_Zavadsky_ENGIE_Impact (1).docx"},
    {"id": "co-15", "name": "Trane_Don_Johnson", "label": "Trane (Don Johnson)", "interviewees": "Don Johnson", "files": "CC_Research_Don_Johnson_Trane_Transcript.docx"},
    {"id": "co-16", "name": "Trane_Phil_Combs", "label": "Trane (Phil Combs)", "interviewees": "Phil Combs", "files": "Clear_Current_Interview_Summary_Phil_Combs_Trane.docx (summary only — no transcript available)"},
]

print(f"\nRunning {len(COMPANIES)} company intelligence cards...\n")

for company in COMPANIES:
    print(f"  → {company['label']}")
    focus_task = (
        f"COMPANY: {company['label']}\n"
        f"INTERVIEWEES: {company['interviewees']}\n"
        f"SOURCE FILES: {company['files']}\n\n"
        f"{COMPANY_CARD_TEMPLATE}"
    )
    try:
        msg = client.messages.create(
            model=MODEL,
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"COMPLETE RESEARCH CORPUS:\n\n{MASTER_CORPUS}",
                        "cache_control": {"type": "ephemeral"},
                    },
                    {"type": "text", "text": focus_task},
                ],
            }],
        )
        card_output = msg.content[0].text
        fname = COMPANY_CARD_DIR / f"{company['id']}_{company['name']}.md"
        fname.write_text(
            f"# {company['label']}\n"
            f"# Generated: {datetime.datetime.now().isoformat()}\n\n---\n\n"
            + card_output,
            encoding="utf-8",
        )
        print(f"     ✓ Saved → {fname.name}  ({len(card_output):,} chars)")
    except Exception as e:
        print(f"     ✗ FAILED: {e}")
    time.sleep(AGENT_DELAY_SECONDS)

# ---------------------------------------------------------------------------
# 9. DECISION BRIEF — SYNTHESIS FOR JOHN AND DAN
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 8 — DECISION BRIEF (synthesis for John & Dan)")
print("="*70)

agent_outputs_text = "\n\n".join([
    f"=== AGENT: {a['name']} | {a['title']} ===\n{all_agent_outputs.get(a['id'], 'NOT RUN')}"
    for a in AGENTS
])

DECISION_BRIEF_TASK = f"""You have just read the complete research corpus — every transcript, every secondary document, every competitive intelligence report. You also have the analysis produced by 13 specialist agents who each went deep on their slice of the research.

Here are their outputs:

{agent_outputs_text}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
YOUR TASK: DECISION BRIEF FOR JOHN REUTER (CEO) AND DAN SCHREIBER (CPO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

John and Dan are walking into a board presentation in 72 hours. They need a single document that tells them exactly what to do next. Not a summary of what the research found. A set of decisions.

## THE 5 THINGS THAT ARE DEFINITIVELY TRUE
The 5 findings so well-evidenced — confirmed by multiple interviews and secondary data — that John and Dan can state them to investors without hedging. Cite at least 2 specific sources per finding. No generic observations.

## THE 5 PRODUCT PRIORITIES — IN ORDER
What to build first through fifth. A sequenced roadmap based on what buyers will actually pay for. For each: the capability, who asked for it (name companies and people), their exact words, and the revenue implication.

## THE 2 CALLS TO MAKE THIS WEEK
Which 2 of the 16 interviewees should John or Dan call back within 7 days. Why those two specifically. What is the exact ask on each call.

## THE 3 COMPETITIVE MOVES TO MAKE NOW
Three specific actions Clear Current can take in the next 90 days to widen its moat. Grounded in specific research evidence. Direct about why the window is open and how long it stays open.

## THE 1 THING THAT COULD KILL THIS
The single biggest risk this research reveals — the assumption the evidence does NOT support, or the finding that directly challenges Clear Current's current strategy. Say it plainly.

## THE SLIDE THAT OPENS THE PRESENTATION
Write the exact text: headline, three supporting bullets, and one verbatim quote — for the opening slide of the investor presentation. This slide should make someone lean forward within 10 seconds.
"""

try:
    # Decision brief synthesizes agent outputs — no raw corpus needed (agents already did that work)
    brief_msg = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": DECISION_BRIEF_TASK}],
    )
    brief_output = brief_msg.content[0].text
    brief_path = OUTPUTS_DIR / "99_decision_brief_john_dan.md"
    brief_path.write_text(
        "# Clear Current Technologies — Decision Brief for John & Dan\n"
        f"# Generated: {datetime.datetime.now().isoformat()}\n\n---\n\n"
        + brief_output,
        encoding="utf-8",
    )
    print(f"✓ Decision brief saved → {brief_path}  ({len(brief_output):,} chars)")
except Exception as e:
    print(f"✗ Decision brief FAILED: {e}")

# ---------------------------------------------------------------------------
# 10. DELIVERABLE BUILDERS
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("STEP 9 — DELIVERABLE BUILDERS (the 6 required + 2 bonus)")
print("="*70)

DELIVERABLES_BUILD_DIR = DELIVERABLES_DIR

DELIVERABLE_SYSTEM = (
    "You are building a final consulting deliverable for Clear Current Technologies — "
    "a seed-stage AI-powered commercial energy management SaaS platform. "
    "The CEO is John Reuter. The CPO is Dan Schreiber. "
    "This deliverable will be presented at a board meeting and used to inform a May 2026 fundraise. "
    "Write for a C-suite audience, not an academic one. "
    "John (CEO) wants high-level insight at every major stage: what is happening in the research and why it matters for Clear Current; verbatim customer quotes in the body with attribution; use interview summaries for orientation and transcripts for exact quotes, flagging summary-only inference. "
    "Be specific, visual where possible using markdown tables and structured lists, and cite real interview evidence. "
    "Every claim must trace back to something a real interviewee said or a real data point from the secondary research. "
    "Do not restate baseline product capabilities the team already knows unless research adds nuance. "
    "Close with a section titled exactly ## Implications for Clear Current (tagged): bullets prefixed with "
    "[FITS CURRENT PRODUCT], [EXTENDS ROADMAP], or [NET NEW / RESEARCH-ONLY]. "
    "Do not claim fully autonomous dispute resolution or closed-loop utility refunds as shipped unless the corpus says so."
)
if PRODUCT_CONTEXT_FOR_AGENTS:
    DELIVERABLE_SYSTEM += (
        "\n\n━━━ CLEAR CURRENT PRODUCT CONTEXT ━━━\n"
        f"{PRODUCT_CONTEXT_FOR_AGENTS}\n"
        "━━━ END PRODUCT CONTEXT ━━━\n"
    )

DELIVERABLES = [
    {
        "id": "d01",
        "filename": "D01_Engagement_Trigger_Map.md",
        "title": "Engagement Trigger Map",
        "task": """Build the complete Engagement Trigger Map for Clear Current Technologies.

This map answers: what causes a corporate energy manager to open the platform?

Structure it as a master table with these columns:
| Trigger | Trigger Type | Vertical | Who Feels It | Urgency | Platform Entry Point | First Action in Platform | Evidence (who said what) |

TRIGGER TYPES: Reactive (problem-driven) | Proactive (planning-driven) | System-Initiated (alert-driven)

VERTICALS: Higher Education | Healthcare | Multi-Site Hospitality | Multi-Site Restaurant | Commercial Real Estate | Energy Services

Include every trigger surfaced across all 16 interviews. Aim for 25–35 rows. Every row must have at least one named interviewee or document as evidence.

After the table, write a TRIGGER FREQUENCY ANALYSIS section:
- Which triggers happen monthly vs. quarterly vs. annually?
- Which triggers produce the highest-urgency platform sessions?
- Which triggers are currently unserved by any competitor?

Then write an ALERT DESIGN IMPLICATIONS section:
- What does this map tell Dan Schreiber about how to design the alert/notification system?
- Which triggers should produce a push notification vs. an email vs. a dashboard flag?"""
    },
    {
        "id": "d02",
        "filename": "D02_Seasonal_Energy_Calendar.md",
        "title": "Seasonal Energy Management Calendar",
        "task": """Build the complete Seasonal Energy Management Calendar for Clear Current Technologies.

This calendar shows how an energy manager's workflow changes across a full year — and where the platform fits at each moment.

Structure it as two layers:

LAYER 1 — MASTER CALENDAR TABLE
One row per month (January through December). Columns:
| Month | What's Happening | Key Deadlines | Dominant Trigger Type | Platform Use Case | Verticals Most Active | Verbatim Evidence |

Use real interview data for every month. If an interviewee described something that happens in a specific month or quarter — put it here.

LAYER 2 — VERTICAL-SPECIFIC CALENDARS
Produce a separate condensed calendar for each vertical:
- Higher Education (Texas State, UT Austin)
- Healthcare (Ascension, Bon Secours, Houston Methodist, Intermountain, AdventHealth)
- Multi-Site Restaurant (Panda Restaurant Group)
- Multi-Site Hospitality (Hyatt)
- Commercial Real Estate (JLL)
- Energy Services / Consulting (ENGIE Impact, Trane)

For each vertical, identify: the 2-3 months where they are most engaged, their biggest annual deadline, and the one platform feature they'd use most in that season.

End with: PRODUCT TIMING IMPLICATIONS — what does this calendar tell Dan about when to run pilots, when to push feature releases, and when NOT to distract energy managers with onboarding?"""
    },
    {
        "id": "d03",
        "filename": "D03_User_Journey_Maps.md",
        "title": "User Journey Maps (6 Core Journeys)",
        "task": """Build all 6 core user journey maps for Clear Current Technologies. These are the primary required deliverable for this engagement.

For EACH of the 6 journeys below, use this exact structure:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# JOURNEY: [Name]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

**Primary Persona:** [name a real interviewee who represents this journey]
**Vertical(s):** [which industries this journey applies to]
**Frequency:** [how often this journey occurs]
**Urgency:** [HIGH / MEDIUM / LOW]

## THE TRIGGER
What specifically causes this journey to begin? Use real interview language. What does the user see, feel, or receive that kicks this off?

## STEP-BY-STEP FLOW
Numbered steps. For each step: what the user does, what question they're trying to answer, what the platform shows them, and what decision they make. Be specific enough that a product designer could wireframe it.

## DECISION POINTS
Where does the user have to make a choice? What happens on each branch?

## THE OUTPUT
What does the user produce at the end of this journey? (A report? An alert response? A budget number? An email to their CFO?)

## WHAT MAKES THIS HARD TODAY
What do interviewees currently do instead of using Clear Current? What breaks down? Quote them directly.

## PLATFORM DESIGN REQUIREMENTS
3–5 specific product requirements this journey reveals. Not wishes — requirements grounded in what interviewees said they need.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

THE 6 JOURNEYS TO MAP:
1. Monthly Bill Review & Validation
2. Anomaly Detection & Response
3. Budgeting & Forecasting
4. Campus / Multi-Site Chargeback & Cost Allocation
5. Regulatory Rate Case & Tariff Change Response
6. Executive & Board Reporting"""
    },
    {
        "id": "d04",
        "filename": "D04_Opportunity_Heatmap.md",
        "title": "Opportunity Heatmap",
        "task": """Build the complete Opportunity Heatmap for Clear Current Technologies.

The heatmap scores every product opportunity on three axes:
- VALUE: financial impact per customer per year ($)
- FREQUENCY: how often this use case occurs (Daily / Weekly / Monthly / Quarterly / Annually)
- STRATEGIC IMPACT: how much winning here differentiates Clear Current vs. named competitors

PART 1 — MASTER HEATMAP TABLE

| Opportunity | Vertical(s) | Value Score (1-5) | Frequency Score (1-5) | Strategic Score (1-5) | TOTAL | Evidence |
|---|---|---|---|---|---|---|

Include every distinct opportunity surfaced across all 16 interviews. Aim for 20–30 rows.
Score each on a 1–5 scale. Total = sum of three scores (max 15). Sort descending by total.
Evidence column: name the interviewee(s) who validated this opportunity and what they said.

PART 2 — TOP 5 OPPORTUNITIES (DEEP DIVE)
For the top 5 scoring opportunities, write a 1-paragraph brief covering:
- The exact dollar value at stake (quantified from interview evidence where possible)
- Which specific companies confirmed this need and what they said
- Why competitors don't solve this today
- What Clear Current needs to build to capture it

PART 3 — COMPETITIVE WHITE SPACE ANALYSIS
Which opportunities have a strategic score of 5 — meaning Clear Current is the only platform positioned to capture them? For each, explain why the window exists and how long it will stay open.

PART 4 — WHAT THIS MEANS FOR PRICING & PACKAGING
Based on the heatmap, what does the research suggest about how to tier and price Clear Current's platform?"""
    },
    {
        "id": "d05",
        "filename": "D05_Product_Module_Recommendations.md",
        "title": "Product Module Recommendations",
        "task": """Build the complete Product Module Recommendations document for Clear Current Technologies.

Produce a full recommendation for each of the 5 core product modules. Use this exact structure for every module:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# MODULE [N]: [Name]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## Problem Definition
What specific, documented problem does this module solve? State it precisely. Cite the interviews that confirm this problem exists. Include the dollar cost of NOT solving it where evidence allows.

## Target User Persona
Not a generic persona. Name the actual people from the research who represent this persona (e.g., "Roger Goldstein at Panda, Sean Sevy at Houston Methodist"). Describe their day-to-day, their constraints, what they care about, and what would make them champion this module internally.

## Business Value
Quantified where possible. What is the ROI for a customer? What evidence from the research supports this number? Be honest about what is documented vs. estimated.

## Frequency of Use
How often does a user engage with this module in a typical month? What drives that frequency?

## Strategic Differentiation
Which named competitors offer this? What do they get wrong or miss entirely? Why does Clear Current's AI-native approach produce a better outcome? Cite specific competitive intelligence from the research.

## MVP Scope
What is the minimum set of capabilities Dan Schreiber needs to ship for this module to be valuable on day one? Be specific enough that an engineer could build a sprint plan from this. What is NOT in the MVP.

## Rollout Sequence
Which vertical do you pilot this with first and why? What does the expansion path look like?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

THE 5 MODULES:
1. Billing Error Audit Engine
2. Anomaly Detection & Alerting
3. Regulatory Intelligence Monitor
4. Tariff Optimization
5. Campus / Multi-Site Chargeback Intelligence"""
    },
    {
        "id": "d06",
        "filename": "D06_Presentation_Narrative.md",
        "title": "Final Presentation Narrative (1-Hour Structure)",
        "task": """Build the complete narrative structure for the Clear Current Technologies final 1-hour presentation to John Reuter and Dan Schreiber.

This is not a slide-by-slide script. It is a narrative architecture — the story arc, the key moments, the evidence that lands hardest, and the exact language to use at each turn.

## ACT 1: THE OPENING (0–5 min) — Make them lean forward
- The one thing that will surprise John and Dan in the first 60 seconds
- The opening statement (write it word for word)
- The first data point that reframes how they see their own market
- The verbatim quote from a real interviewee that opens the room

## ACT 2: WHO WE TALKED TO (5–10 min) — Establish credibility
- How to present the 16 interviews in a way that lands with a board
- The one or two interviewees whose name/company alone signals market validation
- What the diversity of verticals proves about TAM

## ACT 3: WHAT WE FOUND — THE 5 CORE FINDINGS (10–30 min)
For each of the 5 most important findings, write:
- The finding stated as a one-sentence headline
- The two or three pieces of evidence that prove it
- The verbatim quote that makes it real
- The product or GTM implication John and Dan should take away

## ACT 4: THE DELIVERABLES WALK (30–45 min)
How to present the Trigger Map, Seasonal Calendar, Journey Maps, and Opportunity Heatmap in 15 minutes without losing the room. What to emphasize on each, what to skip, what question each slide answers for a C-suite audience.

## ACT 5: PRODUCT MODULE RECOMMENDATIONS (45–55 min)
How to present the 5 module recommendations. Which one to lead with and why. How to frame the sequencing for John and Dan so they understand why order matters.

## ACT 6: THE CLOSE (55–60 min)
- The single call to action for John and Dan
- The one number they should walk out of the room ready to quote to investors
- The closing statement (write it word for word)"""
    },
    {
        "id": "d07",
        "filename": "D07_AI_Chat_Interaction_Flows.md",
        "title": "AI Chat Interaction Flows (Bonus)",
        "task": """Build the sample AI Chat Interaction Flows for Clear Current Technologies — a bonus deliverable from the project scope.

Dan Schreiber's product has three tiers: Chat (React to User), Dashboard (Push to User), Autonomous Agent (Do for User). The chat interface is the primary interaction layer for complex analysis.

For each of the 5 modules, write a realistic sample conversation between an energy manager and the Clear Current AI. Each conversation should:

1. Start with the trigger — what just happened that caused the user to open the chat
2. Show the user's first message (natural language, not technical)
3. Show the AI's response — specific, data-driven, actionable
4. Continue for 4–6 turns until the user has what they need
5. End with the AI offering a proactive next step they didn't ask for

CONVERSATION DESIGN RULES (grounded in what interviewees said):
- Andi Ault said "I don't want to chat with it — I want it to flag something." Design the first message as the AI alerting the user, not the user querying the AI.
- Scott Czubkowski's data trust concerns mean the AI must explain its reasoning. Every finding should include "Here's why I flagged this" not just the conclusion.
- Sean Sevy described finding out about anomalies from the bill weeks later. Show how the AI catches this in real time.

MODULES TO COVER:
1. Billing Error Audit: AI finds a rate misclassification on a 47-location portfolio
2. Anomaly Detection: AI flags a 23% consumption spike at a hospital campus
3. Regulatory Intelligence: AI detects a filed rate case that will affect the user's July budget
4. Campus Chargeback: AI identifies a billing formula error in the district energy chargeback model
5. Executive Brief: User asks the AI to prepare a board-ready energy summary for tomorrow's meeting"""
    },
    {
        "id": "d08",
        "filename": "D08_How_Youll_Use_This_Platform.md",
        "title": "Customer-Facing Platform Guide (Bonus)",
        "task": """Build the customer-facing "How You'll Use This Platform" guide for Clear Current Technologies — a bonus deliverable from the project scope.

This is NOT a feature list or a technical manual. It is a story told to a prospective customer — an energy manager, a VP of Facilities, or a Director of Utilities — showing them exactly what their day looks like with Clear Current vs. without it.

Write it in second person ("You"). Make it vivid. Ground every scenario in something a real interviewee described.

## THE MONDAY MORNING SCENARIO
It's Monday. You have 47 utility bills sitting in your inbox. Walk through what happens with Clear Current vs. what happens today (without it). Be specific about time saved, errors caught, and decisions made.

## THE BUDGET SEASON SCENARIO
It's October. Your CFO wants a preliminary energy budget for the fiscal year by November 15. Walk through how Clear Current compresses what used to be a 3-week manual process. Cite the specific pain points interviewees described during budget season.

## THE ANOMALY SCENARIO
A Thursday at 2pm. You get a push notification. Walk through exactly what the platform shows you, what you do next, and what would have happened if you hadn't had Clear Current (how long would it have taken to find? what would it have cost?).

## THE RATE CASE SCENARIO
A rate case was filed with your state PUC six months ago. Without Clear Current, you find out when the new rate appears on your bill — after your budget is already locked. Show how Clear Current catches this in the public docket and alerts you in time to act.

## THE EXECUTIVE REPORT SCENARIO
Your CEO wants an energy performance summary for the board deck. Before Clear Current: two days, three spreadsheets, a lot of manual formatting. Show what this looks like now.

## WHO THIS IS FOR
End with a one-page "Is This Platform For You?" checklist — the characteristics of an organization that will get the most value from Clear Current, grounded in the profiles of the 16 companies interviewed."""
    },
]


def run_deliverable_builder(deliverable: dict, corpus: str) -> str:
    print(f"\n{'='*60}")
    print(f"BUILDING: {deliverable['title']}")
    print(f"Time: {datetime.datetime.now().strftime('%H:%M:%S')}")
    print(f"{'='*60}")

    try:
        msg = client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            system=DELIVERABLE_SYSTEM,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"COMPLETE RESEARCH CORPUS:\n\n{corpus}",
                        "cache_control": {"type": "ephemeral"},
                    },
                    {
                        "type": "text",
                        "text": (
                            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                            f"YOUR TASK: BUILD DELIVERABLE — {deliverable['title']}\n"
                            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                            f"{deliverable['task']}"
                        ),
                    },
                ],
            }],
        )
        output = msg.content[0].text
        print(f"✓ {deliverable['title']} complete — {len(output):,} chars")
        return output
    except Exception as e:
        print(f"✗ {deliverable['title']} FAILED: {e}")
        return f"FAILED: {e}"


print(f"\nBuilding {len(DELIVERABLES)} deliverables...\n")

for deliverable in DELIVERABLES:
    output = run_deliverable_builder(deliverable, MASTER_CORPUS)
    fpath = DELIVERABLES_BUILD_DIR / deliverable["filename"]
    fpath.write_text(
        f"# {deliverable['title']}\n"
        f"# Clear Current Technologies — CDL MBA Consulting Engagement\n"
        f"# Generated: {datetime.datetime.now().isoformat()}\n\n---\n\n"
        + output,
        encoding="utf-8",
    )
    print(f"  → Saved: {fpath.name}")
    if deliverable != DELIVERABLES[-1]:
        print(f"  Waiting {AGENT_DELAY_SECONDS}s...")
        time.sleep(AGENT_DELAY_SECONDS)

# ---------------------------------------------------------------------------
# 11. FINAL SUMMARY
# ---------------------------------------------------------------------------

print("\n" + "="*70)
print("PIPELINE COMPLETE")
print("="*70)

output_files = list(OUTPUTS_DIR.rglob("*.md")) + list(OUTPUTS_DIR.rglob("*.txt"))
print(f"\nTotal output files generated: {len(output_files)}")
print(f"\nOutput directory: {OUTPUTS_DIR}\n")

for f in sorted(output_files):
    rel = f.relative_to(OUTPUTS_DIR)
    size_kb = f.stat().st_size // 1024
    print(f"  {rel}  ({size_kb} KB)")

print("\n✓ All done. Review outputs in research/outputs/")
print("✓ Company cards → research/outputs/company_cards/")
print("✓ Deliverables  → research/outputs/deliverables/")
print("✓ Decision brief → research/outputs/99_decision_brief_john_dan.md")
