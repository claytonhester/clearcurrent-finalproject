"""
Shared utilities for all three Clear Current research pipelines.
"""
import os, sys, time, datetime
from pathlib import Path
from dotenv import load_dotenv

env_path = Path(__file__).parent.parent / ".env"
load_dotenv(dotenv_path=env_path)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
if not ANTHROPIC_API_KEY:
    print("ERROR: ANTHROPIC_API_KEY not found.")
    print(f"Looked in: {env_path}")
    sys.exit(1)

import anthropic
import docx
import pdfplumber

# ---------------------------------------------------------------------------
# MODELS & SETTINGS
# ---------------------------------------------------------------------------
MODEL       = "claude-sonnet-4-6"
HAIKU_MODEL = "claude-haiku-4-5-20251001"
MAX_TOKENS  = 16000
DELAY       = 120

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

PRODUCT_CONTEXT_FILENAME = "CLEAR_CURRENT_PRODUCT_CONTEXT.md"


def load_product_context(research_dir: Path | None = None) -> str:
    """
    Load canonical Clear Current product/strategy context for agent system prompts.
    research_dir defaults to the directory containing this module (research/).
    """
    rd = research_dir if research_dir is not None else Path(__file__).parent
    path = rd / PRODUCT_CONTEXT_FILENAME
    if not path.exists():
        print(f"⚠  {PRODUCT_CONTEXT_FILENAME} not found at {path} — agents run without product anchor.")
        return ""
    return path.read_text(encoding="utf-8", errors="replace").strip()


# ---------------------------------------------------------------------------
# EXTRACTION
# ---------------------------------------------------------------------------

def extract_docx(path: Path) -> str:
    doc = docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)

def extract_pdf(path: Path) -> str:
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                parts.append(txt)
    return "\n".join(parts)

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":   return extract_docx(path)
    elif suffix == ".pdf":  return extract_pdf(path)
    elif suffix in (".txt", ".md"): return path.read_text(encoding="utf-8", errors="replace")
    return ""

def build_corpus(folders: list[tuple[str, Path]], extra_files: list[Path] = None) -> tuple[str, int]:
    """
    Build a corpus string from a list of (label, folder) tuples.
    Returns (corpus_text, total_chars).
    Prints a summary of what was loaded and warns if near token limit.
    """
    parts = []
    total_chars = 0

    for label, folder in folders:
        if not folder.exists():
            print(f"  ⚠  Folder not found: {folder}")
            continue
        files = sorted([
            f for f in folder.rglob("*.*")
            if f.suffix.lower() in (".docx", ".pdf", ".txt", ".md")
        ])
        print(f"\n[{label}]")
        for fpath in files:
            try:
                text = extract_text(fpath)
                total_chars += len(text)
                parts.append(
                    f"\n\n═══ SOURCE: {fpath.name} | FOLDER: {label} ═══\n"
                    f"{text}\n"
                    f"═══ END: {fpath.name} ═══\n"
                )
                print(f"  ✓ {fpath.name} — {len(text):,} chars")
            except Exception as e:
                print(f"  ✗ {fpath.name} — FAILED: {e}")

    for fpath in (extra_files or []):
        if not fpath.exists():
            print(f"  ⚠  Extra file not found: {fpath}")
            continue
        try:
            text = extract_text(fpath)
            total_chars += len(text)
            parts.append(
                f"\n\n═══ SOURCE: {fpath.name} | FOLDER: extra ═══\n"
                f"{text}\n"
                f"═══ END: {fpath.name} ═══\n"
            )
            print(f"  ✓ {fpath.name} — {len(text):,} chars")
        except Exception as e:
            print(f"  ✗ {fpath.name} — FAILED: {e}")

    corpus = "\n".join(parts)
    approx_tokens = len(corpus) // 4
    print(f"\nTotal: {total_chars:,} chars | ~{approx_tokens:,} estimated tokens")
    if approx_tokens > 185000:
        print(f"⚠  WARNING: corpus is near the 200K token limit — consider trimming secondary sources.")
    return corpus, total_chars

# ---------------------------------------------------------------------------
# CACHED API CALL
# ---------------------------------------------------------------------------

def call_with_cache(
    corpus: str,
    task: str,
    system: str = None,
    model: str = None,
    max_tokens: int = None,
) -> str:
    """
    Send a prompt to Claude with the corpus as a cached content block.
    Returns the response text or an error string.
    """
    m = model or MODEL
    mt = max_tokens or MAX_TOKENS
    kwargs = dict(
        model=m,
        max_tokens=mt,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"COMPLETE RESEARCH CORPUS:\n\n{corpus}",
                    "cache_control": {"type": "ephemeral"},
                },
                {"type": "text", "text": task},
            ],
        }],
    )
    if system:
        kwargs["system"] = system
    msg = client.messages.create(**kwargs)
    usage = msg.usage
    cr = getattr(usage, "cache_read_input_tokens", 0) or 0
    cw = getattr(usage, "cache_creation_input_tokens", 0) or 0
    out = msg.content[0].text
    print(f"  tokens — cache_write={cw:,}  cache_read={cr:,}  output={len(out):,} chars")
    return out

def save_output(path: Path, header: str, content: str) -> None:
    path.write_text(header + content, encoding="utf-8")
    size_kb = path.stat().st_size // 1024
    print(f"  → Saved: {path.name}  ({size_kb} KB)")
